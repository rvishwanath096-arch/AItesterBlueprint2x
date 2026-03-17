import docx
import re
import os

def add_formatted_text(p, text):
    # Split by bold (**text**) or italic (*text*)
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            p.add_run(part[1:-1]).italic = True
        else:
            p.add_run(part)

doc = docx.Document()
# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'

file_path = r'e:\AI Testing\Projects\AItesterBlueprint2x\Project_4_RICEPOT_TestPlanGenarator\VWO_Login_Test_Plan.md'
out_path = r'e:\AI Testing\Projects\AItesterBlueprint2x\Project_4_RICEPOT_TestPlanGenarator\VWO_Login_Test_Plan.docx'

if not os.path.exists(file_path):
    print("Markdown file not found")
    exit(1)

with open(file_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
    if line == '---':
        continue
    
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        add_formatted_text(p, line[2:])
    elif line.startswith('- [ ] '):
        p = doc.add_paragraph(style='List Bullet')
        add_formatted_text(p, "\u2610 " + line[6:])
    else:
        p = doc.add_paragraph()
        add_formatted_text(p, line)

doc.save(out_path)
print(f"Successfully generated {out_path}")
