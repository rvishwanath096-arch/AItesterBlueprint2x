import docx
import os
import sys

def extract_text_from_docx(file_path, output_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    doc = docx.Document(file_path)
    full_text: list[str] = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(full_text))

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python extract_docx.py <file_path> <output_path>")
    else:
        extract_text_from_docx(sys.argv[1], sys.argv[2])
